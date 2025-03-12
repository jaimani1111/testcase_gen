import os
from dotenv import load_dotenv
import pandas as pd
from groq import Groq
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
import streamlit as st
from docx import Document

# Load environment variables
load_dotenv()
GROQ_API_KEY = os.getenv("GROQ_API_KEY")

print("Environment variables loaded successfully.")

# Initialize Groq client
groq_client = Groq(api_key=GROQ_API_KEY)
print("Groq client initialized successfully.")

# Load precomputed Chroma DB
try:
    print("Loading Chroma DB...")
    vector_db = Chroma(persist_directory="./chroma_db", embedding_function=HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2"))
    print("Chroma DB loaded successfully.")
except Exception as e:
    print(f"Error loading Chroma DB: {e}")
    st.error(f"Error loading Chroma DB: {e}")
    st.stop()

# Function to generate new test cases using Groq
def generate_test_case(insurance_type, region, line_of_business, user_requirements):
    query = f"Generate a test case for Insurance Type: {insurance_type}, Region: {region}, Line of Business: {line_of_business}, User Requirements: {user_requirements}. Use the same format as the test bank ie Sl No.,Requirement ID,Test Case ID,Module,LOB,Region,Test Case Description,Execution Steps,Expected Result."
    print(f"Query sent to Groq: {query}")
    
    # Use Groq client to generate a response
    response = groq_client.chat.completions.create(
        messages=[{"role": "user", "content": query}],
        model="llama3-70b-8192",
    )
    print("Response received from Groq.")
    return response.choices[0].message.content

# Streamlit App
st.title("Insurance Test Case Generator")

# Inputs
insurance_type = st.selectbox("Insurance Type", ["Health", "Life", "Auto"])
region = st.selectbox("Region", ["Asia", "Europe", "ANZ", "North America"])
line_of_business = st.selectbox("Line of Business", ["Underwriting", "Claims", "Policy Management"])
user_requirements = st.text_area("User Requirements (BDD/TDD)")

# Generate Test Cases
if st.button("Generate Test Cases"):
    print("Generate Test Cases button clicked.")
    new_test_case = generate_test_case(insurance_type, region, line_of_business, user_requirements)
    print("New test case generated.")
    
    # Assuming the response is in the format you need, you can now parse it into structured data
    # Example of how the test case might be structured:
    test_case_data = {
        "Sl No.": 1,
        "Requirement ID": "REQ001",
        "Test Case ID": "TCNEW",
        "Module": line_of_business,
        "LOB": line_of_business,
        "Region": region,
        "Test Case Description": new_test_case,
        "Execution Steps": "1. Execute step 1.\n2. Execute step 2.",
        "Expected Result": "The system behaves as expected."
    }

    # Display the generated test case in a readable format (no JSON, just the test case)
    st.write("### Generated Test Case")
    st.write(f"**Sl No.**: {test_case_data['Sl No.']}")
    st.write(f"**Requirement ID**: {test_case_data['Requirement ID']}")
    st.write(f"**Test Case ID**: {test_case_data['Test Case ID']}")
    st.write(f"**Module**: {test_case_data['Module']}")
    st.write(f"**LOB**: {test_case_data['LOB']}")
    st.write(f"**Region**: {test_case_data['Region']}")
    st.write(f"**Test Case Description**: {test_case_data['Test Case Description']}")
    st.write(f"**Execution Steps**:\n{test_case_data['Execution Steps']}")
    st.write(f"**Expected Result**:\n{test_case_data['Expected Result']}")
    
    # Save to CSV with the required format
    df = pd.DataFrame([test_case_data])
    csv = df.to_csv(index=False)
    st.download_button("Download as CSV", csv, file_name="test_cases.csv")
    print("Test case saved to CSV.")

    # Save to DOC with the required format
    doc = Document()
    doc.add_paragraph(f"Sl No.: {test_case_data['Sl No.']}")
    doc.add_paragraph(f"Requirement ID: {test_case_data['Requirement ID']}")
    doc.add_paragraph(f"Test Case ID: {test_case_data['Test Case ID']}")
    doc.add_paragraph(f"Module: {test_case_data['Module']}")
    doc.add_paragraph(f"LOB: {test_case_data['LOB']}")
    doc.add_paragraph(f"Region: {test_case_data['Region']}")
    doc.add_paragraph(f"Test Case Description: {test_case_data['Test Case Description']}")
    doc.add_paragraph(f"Execution Steps: {test_case_data['Execution Steps']}")
    doc.add_paragraph(f"Expected Result: {test_case_data['Expected Result']}")
    
    doc.save("test_cases.docx")
    with open("test_cases.docx", "rb") as file:
        st.download_button("Download as DOC", file, file_name="test_cases.docx")
    print("Test case saved to DOC.")
