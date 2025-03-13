import os
import pandas as pd
from dotenv import load_dotenv
from groq import Groq
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
import streamlit as st
from docx import Document
from io import BytesIO

# Load environment variables
load_dotenv()

GROQ_API_KEY = os.getenv("GROQ_API_KEY")

groq_client = Groq(api_key=GROQ_API_KEY)

embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

# Load test bank
test_bank = pd.read_csv("test.csv")

# Load Chroma DB with precomputed embeddings
try:
    vector_db = Chroma(persist_directory="./chroma_db", embedding_function=embeddings)
except Exception as e:
    st.error(f"Error loading Chroma DB: {e}")
    st.stop()

# Function to clean description
def clean_description(description):
    return description.replace("*", "").strip()

# Function to generate new test cases
def generate_test_case(insurance_type, region, line_of_business, user_requirements):
    query = f"Insurance Type: {insurance_type}, Region: {region}, Line of Business: {line_of_business}, User Requirements: {user_requirements}"
    similar_cases = vector_db.similarity_search(query, k=3)
    context = "\n".join([case.page_content for case in similar_cases])
    
    prompt = f"""
    Generate a completely new test case for Insurance Type: {insurance_type}, 
    Region: {region}, 
    Line of Business: {line_of_business}, 
    User Requirements: {user_requirements}. 

    Ensure that the test case is different from the provided cases in: {context}

    Format:
    - Sl No.
    - Requirement ID
    - Test Case ID
    - Module
    - LOB
    - Region
    - Test Case Description
    - Execution Steps
    - Expected Result
    """

    response = groq_client.chat.completions.create(
        messages=[{"role": "user", "content": prompt}],
        model="llama3-70b-8192",
    )

    return response.choices[0].message.content

# Streamlit UI
st.markdown("<h1 style='text-align: center;'>XC AI Assisted Test Cases Generator</h1>", unsafe_allow_html=True)

with st.sidebar:
    st.header("Select Inputs")
    insurance_type = st.selectbox("Insurance Type", ["Select", "Health", "Auto", "Home", "Life"])
    region = st.selectbox("Region", ["Select", "North America", "Europe", "Asia Pacific", "ANZ"])
    line_of_business = st.selectbox("Line of Business", ["Select", "Retail", "Commercial", "Enterprise"])

    st.subheader("Upload Acceptance Criteria File")
    criteria_file = st.file_uploader("Drop a file here or browse", type=["txt", "csv", "docx"])

    def extract_criteria(file):
        if file is None:
            return ""
        if file.type == "text/plain":
            return file.getvalue().decode("utf-8")
        if file.type == "text/csv":
            df = pd.read_csv(file)
            return "\n".join(df.iloc[:, 0].dropna().astype(str))
        if file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = Document(file)
            return "\n".join([p.text for p in doc.paragraphs])
        return ""

    uploaded_criteria = extract_criteria(criteria_file)

    st.subheader("Enter Acceptance Criteria")
    if criteria_file:
        st.text_area("Acceptance Criteria", uploaded_criteria, height=100, disabled=True)
    else:
        manual_criteria = st.text_area("Write Acceptance Criteria", "", height=100)

    if st.button("Generate Test Case"):
        st.session_state.generate_clicked = True
    else:
        st.session_state.generate_clicked = False

final_criteria = uploaded_criteria if uploaded_criteria else manual_criteria

if final_criteria and st.session_state.generate_clicked:
    st.success("Acceptance Criteria Ready.")
    
    st.subheader("Generated Test Cases")
    
    new_test_case = generate_test_case(insurance_type, region, line_of_business, final_criteria)

    cleaned_description = clean_description(new_test_case)

    new_test_case_data = {
        "Sl No.": len(test_bank) + 1,
        "Requirement ID": "REQ_NEW",
        "Test Case ID": f"TC_{len(test_bank) + 1}",
        "Module": line_of_business,
        "LOB": insurance_type,
        "Region": region,
        "Test Case Description": cleaned_description,
        "Execution Steps": "Step 1: Navigate to the module.\nStep 2: Enter details.\nStep 3: Submit and verify.",
        "Expected Result": "Successfully processed.",
    }

    new_test_case_df = pd.DataFrame([new_test_case_data])
    updated_test_bank = pd.concat([test_bank, new_test_case_df], ignore_index=True)
    updated_test_bank.to_csv("test.csv", index=False)

    st.success("Test case saved to CSV successfully.")

    # Display test case in a tabular format
    st.write("### Generated Test Cases in Table Format")
    st.dataframe(new_test_case_df)

    st.subheader("Download Test Cases")
    file_format = st.radio("Select file format:", ["CSV", "Word"])

    def convert_to_csv(test_case_data):
        df = pd.DataFrame([test_case_data])
        df["Execution Steps"] = df["Execution Steps"].str.replace("\n", "\\n")
        output = BytesIO()
        df.to_csv(output, index=False)
        output.seek(0)
        return output

    def convert_to_word(test_case_data):
        doc = Document()
        doc.add_heading("Generated Test Case", level=1)
        doc.add_paragraph(f"Sl No.: {test_case_data['Sl No.']}")
        doc.add_paragraph(f"Requirement ID: {test_case_data['Requirement ID']}")
        doc.add_paragraph(f"Test Case ID: {test_case_data['Test Case ID']}")
        doc.add_paragraph(f"Module: {test_case_data['Module']}")
        doc.add_paragraph(f"LOB: {test_case_data['LOB']}")
        doc.add_paragraph(f"Region: {test_case_data['Region']}")
        doc.add_paragraph(f"Test Case Description: {test_case_data['Test Case Description']}")
        doc.add_paragraph(f"Execution Steps: {test_case_data['Execution Steps']}")
        doc.add_paragraph(f"Expected Result: {test_case_data['Expected Result']}")
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output

    if file_format == "CSV":
        st.download_button("Download CSV", convert_to_csv(new_test_case_data), "test_case.csv", mime="text/csv")
    elif file_format == "Word":
        st.download_button("Download Word", convert_to_word(new_test_case_data), "test_case.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
